"""
Clean Quote - Simplifié pour workflow ancien devis
Vide UNIQUEMENT les tableaux, garde tout le reste (nom, tel, email customer success)
"""

import json
import logging
import io
from datetime import datetime
from typing import Dict, Any
import azure.functions as func
from docx import Document

import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))
from shared.blob_client import get_blob_client
from shared.validators import validate_required_fields
from shared.logger import setup_logger
from shared.config import CONTAINER_TEMPLATES, get_user_file_path

logger = setup_logger(__name__)


def empty_table_rows(doc: Document) -> int:
    """
    Vide toutes les lignes de données des tableaux (garde les headers)

    Args:
        doc: Document Word

    Returns:
        Nombre de lignes vidées
    """
    rows_emptied = 0

    for table in doc.tables:
        # Garder le header (première ligne)
        # Vider toutes les autres lignes
        if len(table.rows) > 1:
            # Commencer à l'index 1 pour garder le header
            for row_idx in range(len(table.rows) - 1, 0, -1):  # Parcourir à l'envers pour éviter problèmes d'index
                for cell in table.rows[row_idx].cells:
                    cell.text = ""  # Vider la cellule
                rows_emptied += 1

    logger.info(f"Emptied {rows_emptied} table rows across {len(doc.tables)} tables")
    return rows_emptied


def clean_quote(req: func.HttpRequest) -> func.HttpResponse:
    """
    Nettoie un ancien devis en vidant UNIQUEMENT les tableaux

    Workflow:
    1. Charge le fichier depuis Blob Storage (users/{displayName}/ancien_devis.docx)
    2. Vide les lignes de données des tableaux (garde headers)
    3. Sauvegarde le fichier nettoyé dans users/{displayName}/temp_working.docx
    4. Retourne le chemin du fichier de travail

    Request body:
    {
        "blob_path": "users/Jean Dupont/ancien_devis.docx",
        "user_folder": "Jean Dupont"
    }

    Response:
    {
        "success": true,
        "working_file": "users/Jean Dupont/temp_working.docx",
        "blob_url": "https://...",
        "tables_found": 3,
        "rows_emptied": 15
    }
    """
    logger.info("Clean quote endpoint called")

    try:
        # Parse request body
        try:
            req_body = req.get_json()
        except ValueError:
            return func.HttpResponse(
                json.dumps({"error": "Invalid JSON in request body"}),
                status_code=400,
                mimetype="application/json"
            )

        # Validate required fields
        if not validate_required_fields(req_body, ["blob_path", "user_folder"]):
            return func.HttpResponse(
                json.dumps({
                    "error": "Missing required fields",
                    "required": ["blob_path", "user_folder"]
                }),
                status_code=400,
                mimetype="application/json"
            )

        blob_path = req_body["blob_path"]
        user_folder = req_body["user_folder"]

        logger.info(f"Cleaning quote from: {blob_path}")

        # Initialize Blob client
        blob_client = get_blob_client()

        # Container: word-templates (contient general/ et dossiers utilisateurs)
        container_name = CONTAINER_TEMPLATES

        # Télécharger le fichier depuis Blob Storage
        try:
            file_bytes = blob_client.download_blob(container_name, blob_path)
            logger.info(f"Downloaded {len(file_bytes)} bytes from blob")
        except Exception as e:
            logger.error(f"Failed to download blob: {str(e)}")
            return func.HttpResponse(
                json.dumps({
                    "error": "File not found in Blob Storage",
                    "blob_path": blob_path
                }),
                status_code=404,
                mimetype="application/json"
            )

        # Charger le document Word
        doc = Document(io.BytesIO(file_bytes))
        tables_count = len(doc.tables)

        # Vider les tableaux
        rows_emptied = empty_table_rows(doc)

        # Sauvegarder le document nettoyé
        output_bytes_io = io.BytesIO()
        doc.save(output_bytes_io)
        output_bytes = output_bytes_io.getvalue()

        # Upload vers Blob Storage comme fichier de travail
        working_file_path = get_user_file_path(user_folder, "temp_working.docx")

        blob_url = blob_client.upload_blob(
            container_name=container_name,
            blob_name=working_file_path,
            data=output_bytes,
            overwrite=True
        )

        logger.info(f"Working file saved to: {working_file_path}")

        # Retourner résultat
        response_data = {
            "success": True,
            "working_file": working_file_path,
            "blob_url": blob_url,
            "tables_found": tables_count,
            "rows_emptied": rows_emptied
        }

        return func.HttpResponse(
            json.dumps(response_data),
            status_code=200,
            mimetype="application/json"
        )

    except Exception as e:
        logger.error(f"Error cleaning quote: {str(e)}", exc_info=True)
        return func.HttpResponse(
            json.dumps({
                "error": "Failed to clean quote",
                "message": str(e)
            }),
            status_code=500,
            mimetype="application/json"
        )
