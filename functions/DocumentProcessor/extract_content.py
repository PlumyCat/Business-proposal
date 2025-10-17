"""
Extract Content Handler
Extrait le contenu textuel et les tableaux d'un document Word (.docx)
"""

import json
import logging
import io
from typing import Dict, List, Any
import azure.functions as func
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

import sys
sys.path.append('..')
from shared.blob_client import get_blob_client
from shared.validators import validate_required_fields, ValidationError
from shared.logger import setup_logger

logger = setup_logger(__name__)


def extract_tables_from_document(doc: Document) -> List[Dict[str, Any]]:
    """
    Extract all tables from document

    Args:
        doc: python-docx Document object

    Returns:
        List of table dictionaries with headers and rows
    """
    tables_data = []

    for idx, table in enumerate(doc.tables):
        try:
            # Extract table data
            table_dict = {
                "table_index": idx,
                "rows": []
            }

            # First row is usually headers
            if table.rows:
                first_row = table.rows[0]
                headers = [cell.text.strip() for cell in first_row.cells]
                table_dict["headers"] = headers

                # Extract data rows
                for row in table.rows[1:]:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_dict["rows"].append(row_data)

            tables_data.append(table_dict)

        except Exception as e:
            logger.warning(f"Error extracting table {idx}: {str(e)}")
            continue

    logger.info(f"Extracted {len(tables_data)} tables")
    return tables_data


def extract_paragraphs_from_document(doc: Document) -> List[Dict[str, str]]:
    """
    Extract all paragraphs with their styles

    Args:
        doc: python-docx Document object

    Returns:
        List of paragraph dictionaries
    """
    paragraphs_data = []

    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip():  # Skip empty paragraphs
            para_dict = {
                "index": idx,
                "text": para.text.strip(),
                "style": para.style.name if para.style else "Normal"
            }
            paragraphs_data.append(para_dict)

    logger.info(f"Extracted {len(paragraphs_data)} paragraphs")
    return paragraphs_data


def extract_document_structure(doc: Document) -> Dict[str, Any]:
    """
    Extract full document structure (paragraphs + tables)

    Args:
        doc: python-docx Document object

    Returns:
        Dictionary with document structure
    """
    structure = {
        "elements": []
    }

    for element in doc.element.body:
        if isinstance(element, Paragraph):
            structure["elements"].append({
                "type": "paragraph",
                "text": element.text.strip()
            })
        elif hasattr(element, 'tag') and 'tbl' in element.tag:
            structure["elements"].append({
                "type": "table",
                "index": len([e for e in structure["elements"] if e["type"] == "table"])
            })

    return structure


def extract_content(req: func.HttpRequest) -> func.HttpResponse:
    """
    Extract content from a Word document stored in Blob Storage

    Request body:
    {
        "file_id": "unique-file-id"
    }

    Response:
    {
        "file_id": "unique-file-id",
        "text": "Full document text",
        "paragraphs": [...],
        "tables": [...],
        "structure": {...}
    }
    """
    logger.info("Extract content endpoint called")

    try:
        # Parse request body
        try:
            req_body = req.get_json()
        except ValueError:
            return func.HttpResponse(
                json.dumps({"error": "Invalid JSON in request body"}),
                status_code=400,
                mimetype="application/json"
            )

        # Validate required fields
        try:
            validate_required_fields(req_body, ["file_id"])
        except ValidationError as e:
            return func.HttpResponse(
                json.dumps({"error": str(e)}),
                status_code=400,
                mimetype="application/json"
            )

        file_id = req_body["file_id"]
        logger.info(f"Extracting content from file_id: {file_id}")

        # Get blob client and download document
        blob_client = get_blob_client()
        blob_name = f"{file_id}.docx"

        try:
            # Download from devis-sources container
            document_bytes = blob_client.download_blob(
                container_name=blob_client.container_devis,
                blob_name=blob_name
            )
        except Exception as e:
            logger.error(f"Failed to download blob {blob_name}: {str(e)}")
            return func.HttpResponse(
                json.dumps({"error": f"File not found: {file_id}"}),
                status_code=404,
                mimetype="application/json"
            )

        # Parse document with python-docx
        try:
            doc = Document(io.BytesIO(document_bytes))
        except Exception as e:
            logger.error(f"Failed to parse document: {str(e)}")
            return func.HttpResponse(
                json.dumps({"error": "Failed to parse document. Ensure it's a valid .docx file"}),
                status_code=400,
                mimetype="application/json"
            )

        # Extract content
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        paragraphs = extract_paragraphs_from_document(doc)
        tables = extract_tables_from_document(doc)
        structure = extract_document_structure(doc)

        # Build response
        response_data = {
            "file_id": file_id,
            "text": full_text,
            "paragraphs": paragraphs,
            "tables": tables,
            "structure": structure,
            "stats": {
                "total_paragraphs": len(paragraphs),
                "total_tables": len(tables),
                "total_characters": len(full_text)
            }
        }

        logger.info(f"Successfully extracted content from {file_id}")
        return func.HttpResponse(
            json.dumps(response_data),
            status_code=200,
            mimetype="application/json"
        )

    except Exception as e:
        logger.error(f"Unexpected error in extract_content: {str(e)}")
        return func.HttpResponse(
            json.dumps({
                "error": "Internal server error",
                "message": str(e)
            }),
            status_code=500,
            mimetype="application/json"
        )
