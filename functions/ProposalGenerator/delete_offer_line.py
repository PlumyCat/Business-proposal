"""
Delete Offer Line - Supprime une ligne d'offre du document de travail
"""

import json
import logging
import io
import os
from typing import Optional
import azure.functions as func
from docx import Document

import sys
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))
from shared.blob_client import get_blob_client
from shared.validators import validate_required_fields
from shared.logger import setup_logger
from shared.config import CONTAINER_TEMPLATES, get_user_file_path

logger = setup_logger(__name__)

# Mapping service codes Dataverse -> Service display names
SERVICE_CODE_TO_NAME = {
    "480810003": "Service cloud",
    "480810000": "Services Cloud - Téléphonie",
    "480810004": "Supports"
}


def find_table_by_title(doc: Document, title: str) -> Optional[int]:
    """
    Find table index by looking for a paragraph containing the title text
    just before the table.

    Returns table index if found, None otherwise.
    """
    for i, paragraph in enumerate(doc.paragraphs):
        if title.lower() in paragraph.text.lower():
            for table_idx, table in enumerate(doc.tables):
                try:
                    para_element = paragraph._element
                    table_element = table._element

                    parent = para_element.getparent()
                    if parent is not None:
                        children = list(parent)
                        try:
                            para_pos = children.index(para_element)
                            table_pos = children.index(table_element)
                            if table_pos > para_pos:
                                return table_idx
                        except ValueError:
                            continue
                except:
                    continue

    return None


def find_title_paragraph(doc: Document, title: str) -> Optional[int]:
    """
    Find the paragraph index containing the title.

    Returns paragraph index if found, None otherwise.
    """
    for i, paragraph in enumerate(doc.paragraphs):
        if title.lower() in paragraph.text.lower():
            return i
    return None


def delete_table_and_title(doc: Document, table_idx: int, title: str) -> None:
    """
    Delete a table and its title paragraph from the document.
    """
    # Delete the table
    tbl = doc.tables[table_idx]._element
    tbl.getparent().remove(tbl)

    # Find and delete the title paragraph
    para_idx = find_title_paragraph(doc, title)
    if para_idx is not None:
        p = doc.paragraphs[para_idx]._element
        p.getparent().remove(p)
        logger.info(f"Deleted title paragraph '{title}'")


def update_table_total(table: object) -> None:
    """
    Calculate and update the total HT for a service table.
    """
    total_ht = 0.0

    # Sum all rows except header (0) and total (last)
    for row_idx in range(1, len(table.rows) - 1):
        try:
            prix_total_text = table.rows[row_idx].cells[4].text.strip()
            prix_total_text = prix_total_text.replace('€', '').replace(',', '.').strip()
            prix_total = float(prix_total_text)
            total_ht += prix_total
        except (ValueError, IndexError) as e:
            logger.warning(f"Could not parse price in row {row_idx}: {e}")
            continue

    # Update total row
    last_row = table.rows[-1]
    last_row.cells[4].text = f"{total_ht:.2f} €"

    logger.info(f"Updated table total: {total_ht:.2f} €")
    return total_ht


def delete_offer_line(req: func.HttpRequest) -> func.HttpResponse:
    """
    Supprime une ligne d'offre du document de travail

    **Comportement :**
    - Supprime une ligne spécifique identifiée par nom d'offre ou index
    - Met à jour automatiquement le total HT du tableau
    - Si le tableau devient vide (seulement header + total), supprime le tableau et son titre

    Request body:
    {
        "user_folder": "Eric FER",
        "crb02_service": "480810003",
        "offer_name": "Microsoft 365 Apps for Business"  // OU
        "row_index": 2  // Index de la ligne (1 = première ligne après header)
    }

    Response:
    {
        "success": true,
        "working_file": "Eric FER/temp_working.docx",
        "service_name": "Service cloud",
        "deleted_offer": "Microsoft 365 Apps for Business",
        "table_deleted": false,
        "rows_remaining": 3,
        "table_total_ht": "42.00 €"
    }
    """
    logger.info("Delete offer line endpoint called")

    try:
        # Parse request body
        try:
            req_body = req.get_json()
        except ValueError:
            return func.HttpResponse(
                json.dumps({"error": "Invalid JSON in request body"}),
                status_code=400,
                mimetype="application/json"
            )

        # Validate required fields
        if not validate_required_fields(req_body, ["user_folder", "crb02_service"]):
            return func.HttpResponse(
                json.dumps({
                    "error": "Missing required fields",
                    "required": ["user_folder", "crb02_service"]
                }),
                status_code=400,
                mimetype="application/json"
            )

        user_folder = req_body["user_folder"]
        service_code = str(req_body["crb02_service"])

        # Get offer name or row index
        offer_name = req_body.get("offer_name")
        row_index = req_body.get("row_index")

        if not offer_name and row_index is None:
            return func.HttpResponse(
                json.dumps({
                    "error": "Must provide either 'offer_name' or 'row_index'"
                }),
                status_code=400,
                mimetype="application/json"
            )

        # Map service code to name
        service_name = SERVICE_CODE_TO_NAME.get(service_code)
        if not service_name:
            return func.HttpResponse(
                json.dumps({
                    "error": "Invalid or unknown service code",
                    "service_code": service_code,
                    "valid_codes": list(SERVICE_CODE_TO_NAME.keys())
                }),
                status_code=400,
                mimetype="application/json"
            )

        logger.info(f"Deleting offer line from {user_folder}, service: {service_name}")

        # Initialize Blob client
        blob_client = get_blob_client()
        container_name = CONTAINER_TEMPLATES

        # Load working file
        working_file_path = get_user_file_path(user_folder, "temp_working.docx")

        try:
            working_bytes = blob_client.download_blob(container_name, working_file_path)
            logger.info(f"Downloaded working file: {working_file_path}")
        except Exception as e:
            logger.error(f"Failed to download working file: {str(e)}")
            return func.HttpResponse(
                json.dumps({
                    "error": "Working file not found",
                    "working_file": working_file_path,
                    "hint": "Call /api/proposal/prepare-template first"
                }),
                status_code=404,
                mimetype="application/json"
            )

        # Load document
        doc = Document(io.BytesIO(working_bytes))

        # Find table for this service
        table_idx = find_table_by_title(doc, service_name)

        if table_idx is None:
            return func.HttpResponse(
                json.dumps({
                    "error": f"Table for service '{service_name}' not found",
                    "service_name": service_name
                }),
                status_code=404,
                mimetype="application/json"
            )

        table = doc.tables[table_idx]
        logger.info(f"Found table '{service_name}' at index {table_idx}")

        # Find the row to delete
        row_to_delete_idx = None
        deleted_offer_name = None

        if row_index is not None:
            # Use provided row index (1-based, skip header)
            row_to_delete_idx = row_index
            if row_to_delete_idx < 1 or row_to_delete_idx >= len(table.rows) - 1:
                return func.HttpResponse(
                    json.dumps({
                        "error": "Invalid row_index",
                        "row_index": row_index,
                        "valid_range": f"1 to {len(table.rows) - 2}"
                    }),
                    status_code=400,
                    mimetype="application/json"
                )
            deleted_offer_name = table.rows[row_to_delete_idx].cells[0].text

        elif offer_name:
            # Search for offer by name (Désignation column)
            for idx in range(1, len(table.rows) - 1):  # Skip header and total
                cell_text = table.rows[idx].cells[0].text.strip()
                if offer_name.lower() in cell_text.lower():
                    row_to_delete_idx = idx
                    deleted_offer_name = cell_text
                    break

            if row_to_delete_idx is None:
                return func.HttpResponse(
                    json.dumps({
                        "error": f"Offer '{offer_name}' not found in table '{service_name}'"
                    }),
                    status_code=404,
                    mimetype="application/json"
                )

        # Delete the row
        tbl = table._tbl
        tbl.remove(table.rows[row_to_delete_idx]._element)
        logger.info(f"Deleted row {row_to_delete_idx}: {deleted_offer_name}")

        # Check if table is now empty (only header + total remain)
        table_deleted = False
        rows_remaining = len(table.rows) - 2  # Exclude header and total

        if rows_remaining == 0:
            # Delete the entire table and its title
            delete_table_and_title(doc, table_idx, service_name)
            logger.info(f"Table '{service_name}' was empty, deleted table and title")
            table_deleted = True
            table_total_ht = "0.00 €"
        else:
            # Update table total
            table_total_ht = update_table_total(table)
            table_total_ht = f"{table_total_ht:.2f} €"

        # Save document
        output_bytes_io = io.BytesIO()
        doc.save(output_bytes_io)
        output_bytes = output_bytes_io.getvalue()

        # Upload to Blob Storage
        blob_url = blob_client.upload_blob(
            container_name=container_name,
            blob_name=working_file_path,
            data=output_bytes,
            overwrite=True
        )

        logger.info(f"Working file updated: {working_file_path}")

        # Return success response
        response_data = {
            "success": True,
            "working_file": working_file_path,
            "service_name": service_name,
            "deleted_offer": deleted_offer_name,
            "table_deleted": table_deleted,
            "rows_remaining": rows_remaining,
            "table_total_ht": table_total_ht
        }

        return func.HttpResponse(
            json.dumps(response_data),
            status_code=200,
            mimetype="application/json"
        )

    except Exception as e:
        logger.error(f"Error deleting offer line: {str(e)}", exc_info=True)
        return func.HttpResponse(
            json.dumps({
                "error": "Failed to delete offer line",
                "message": str(e)
            }),
            status_code=500,
            mimetype="application/json"
        )
