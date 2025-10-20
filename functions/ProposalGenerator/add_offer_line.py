"""
Add Offer Line - Ajoute une ligne d'offre dans le document de travail
"""

import json
import logging
import io
import os
from typing import Dict, Optional, Tuple
import azure.functions as func
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))
from shared.blob_client import get_blob_client
from shared.validators import validate_required_fields
from shared.logger import setup_logger
from shared.config import CONTAINER_TEMPLATES, get_user_file_path

logger = setup_logger(__name__)

# Mapping service codes Dataverse -> Service display names
SERVICE_CODE_TO_NAME = {
    "480810003": "Service cloud",
    "480810000": "Services Cloud - Téléphonie",
    "480810004": "Supports"
}


def find_table_by_title(doc: Document, title: str) -> Optional[int]:
    """
    Find table index by looking for a paragraph containing the title text
    just before the table.

    Returns table index if found, None otherwise.
    """
    for i, paragraph in enumerate(doc.paragraphs):
        # Check if this paragraph contains the title
        if title.lower() in paragraph.text.lower():
            # Look for the next table after this paragraph
            for table_idx, table in enumerate(doc.tables):
                try:
                    # Get the position of the paragraph
                    para_element = paragraph._element
                    table_element = table._element

                    # Check if table comes after paragraph
                    parent = para_element.getparent()
                    if parent is not None:
                        children = list(parent)
                        try:
                            para_pos = children.index(para_element)
                            table_pos = children.index(table_element)
                            if table_pos > para_pos:
                                return table_idx
                        except ValueError:
                            continue
                except:
                    continue

    return None


def create_service_table(doc: Document, service_name: str) -> int:
    """
    Create a new table for a service with headers and total row.

    Structure:
    - Title paragraph (service name)
    - Table with columns: Désignation | Description | Qté | Prix unitaire | Prix total
    - Header row
    - Total row (initially 0.00 €)

    Returns the index of the created table.
    """
    # Add service title paragraph
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(service_name)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Create table with 5 columns, 2 rows (header + total)
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Light Grid Accent 1'  # Apply a Word style

    # Set column headers
    headers = ['Désignation', 'Description', 'Qté', 'Prix unitaire', 'Prix total']
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        # Make header bold
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Set total row (row 1)
    total_cells = table.rows[1].cells
    # Merge first 4 cells for "Total HT" label
    merged_cell = total_cells[0].merge(total_cells[3])
    merged_cell.text = "Total HT"
    for paragraph in merged_cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in paragraph.runs:
            run.font.bold = True

    # Set initial total to 0.00 €
    total_cells[4].text = "0.00 €"
    for paragraph in total_cells[4].paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in paragraph.runs:
            run.font.bold = True

    # Add a blank paragraph after table for spacing
    doc.add_paragraph()

    # Return the index of the newly created table
    return len(doc.tables) - 1


def find_or_create_service_table(doc: Document, service_name: str) -> Tuple[object, int]:
    """
    Find an existing table for a service or create a new one.

    Returns (table object, table index)
    """
    # Try to find existing table
    table_idx = find_table_by_title(doc, service_name)

    if table_idx is not None:
        logger.info(f"Found existing table for service '{service_name}' at index {table_idx}")
        return doc.tables[table_idx], table_idx

    # Table doesn't exist, create it
    logger.info(f"Creating new table for service '{service_name}'")
    table_idx = create_service_table(doc, service_name)
    return doc.tables[table_idx], table_idx


def update_table_total(table: object) -> None:
    """
    Calculate and update the total HT for a service table.

    Assumes:
    - Row 0 is headers
    - Last row is the total row
    - Column 4 contains the price totals
    """
    total_ht = 0.0

    # Sum all rows except header (0) and total (last)
    for row_idx in range(1, len(table.rows) - 1):
        try:
            # Get prix total from column 4 (index 4)
            prix_total_text = table.rows[row_idx].cells[4].text.strip()
            # Remove € symbol and convert to float
            prix_total_text = prix_total_text.replace('€', '').replace(',', '.').strip()
            prix_total = float(prix_total_text)
            total_ht += prix_total
        except (ValueError, IndexError) as e:
            logger.warning(f"Could not parse price in row {row_idx}: {e}")
            continue

    # Update total row (last row, column 4)
    last_row = table.rows[-1]
    last_row.cells[4].text = f"{total_ht:.2f} €"
    for paragraph in last_row.cells[4].paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in paragraph.runs:
            run.font.bold = True

    logger.info(f"Updated table total: {total_ht:.2f} €")


def add_offer_line(req: func.HttpRequest) -> func.HttpResponse:
    """
    Ajoute une ligne d'offre dans le document de travail

    Le bot lit les offres depuis Dataverse (table crb02_offrebecloud)
    et appelle cet endpoint pour chaque offre à ajouter, une par une.

    **Comportement dynamique :**
    - Si le tableau pour le service n'existe pas → il est créé automatiquement
    - Si le tableau existe déjà → la ligne est ajoutée dedans
    - Chaque tableau a un total HT automatiquement calculé et mis à jour

    **Services supportés :**
    - Service 480810003 (MW) → tableau "Service cloud"
    - Service 480810000 (BUBA) → tableau "Services Cloud - Téléphonie"
    - Service 480810004 (Support) → tableau "Supports"

    **Structure des tableaux créés :**
    - Titre : Nom du service (ex: "Service cloud")
    - Colonnes : Désignation | Description | Qté | Prix unitaire | Prix total
    - Ligne de total HT en bas du tableau

    Request body:
    {
        "user_folder": "Eric FER",
        "offer": {
            "crb02_offrebecloud1": "Microsoft 365 Apps for Business",
            "crb02_description": "Pack Bureautique installable et online – OneDrive 1To",
            "crb02_prixht": 10.00,
            "crb02_service": "480810003",
            "quantity": 3
        }
    }

    Response:
    {
        "success": true,
        "working_file": "Eric FER/temp_working.docx",
        "service_name": "Service cloud",
        "table_created_or_found": "created",
        "rows_in_table": 4,
        "table_total_ht": "30.00 €"
    }
    """
    logger.info("Add offer line endpoint called")

    try:
        # Parse request body
        try:
            req_body = req.get_json()
        except ValueError:
            return func.HttpResponse(
                json.dumps({"error": "Invalid JSON in request body"}),
                status_code=400,
                mimetype="application/json"
            )

        # Validate required fields
        if not validate_required_fields(req_body, ["user_folder", "offer"]):
            return func.HttpResponse(
                json.dumps({
                    "error": "Missing required fields",
                    "required": ["user_folder", "offer"]
                }),
                status_code=400,
                mimetype="application/json"
            )

        user_folder = req_body["user_folder"]
        offer = req_body["offer"]

        # Validate offer fields - use Dataverse field names
        required_offer_fields = ["crb02_offrebecloud1", "crb02_prixht", "crb02_service", "quantity"]
        if not validate_required_fields(offer, required_offer_fields):
            return func.HttpResponse(
                json.dumps({
                    "error": "Missing required offer fields",
                    "required": required_offer_fields
                }),
                status_code=400,
                mimetype="application/json"
            )

        offer_name = offer.get('crb02_offrebecloud1', 'Unknown')
        logger.info(f"Adding offer line to {user_folder}: {offer_name}")

        # Get service code and map to service name
        service_code = str(offer.get('crb02_service', ''))
        service_name = SERVICE_CODE_TO_NAME.get(service_code)

        if not service_name:
            return func.HttpResponse(
                json.dumps({
                    "error": "Invalid or unknown service code",
                    "service_code": service_code,
                    "valid_codes": list(SERVICE_CODE_TO_NAME.keys())
                }),
                status_code=400,
                mimetype="application/json"
            )

        logger.info(f"Service code {service_code} mapped to: {service_name}")

        # Initialize Blob client
        blob_client = get_blob_client()
        container_name = CONTAINER_TEMPLATES

        # Load working file
        working_file_path = get_user_file_path(user_folder, "temp_working.docx")

        try:
            working_bytes = blob_client.download_blob(container_name, working_file_path)
            logger.info(f"Downloaded working file: {working_file_path}")
        except Exception as e:
            logger.error(f"Failed to download working file: {str(e)}")
            return func.HttpResponse(
                json.dumps({
                    "error": "Working file not found",
                    "working_file": working_file_path,
                    "hint": "Call /api/proposal/prepare-template first"
                }),
                status_code=404,
                mimetype="application/json"
            )

        # Load document
        doc = Document(io.BytesIO(working_bytes))

        # Find or create table for this service
        table, table_idx = find_or_create_service_table(doc, service_name)
        logger.info(f"Using table '{service_name}' at index {table_idx}")

        # Extract offer data from Dataverse fields
        designation = offer.get("crb02_offrebecloud1", "")
        description = offer.get("crb02_description", "")
        quantity = offer.get("quantity", 1)
        unit_price = float(offer.get("crb02_prixht", 0))
        total_price = quantity * unit_price

        # Insert new row BEFORE the total row (last row)
        # We need to insert at position len(table.rows) - 1
        new_row_idx = len(table.rows) - 1

        # Get the table element
        tbl = table._tbl
        # Create a new row by copying the structure of an existing row
        new_tr = tbl.tr_lst[0].copy()  # Copy header row structure
        # Insert it before the last row (total row)
        tbl.insert(new_row_idx, new_tr)

        # Now get the newly inserted row from the table
        row = table.rows[new_row_idx]

        # Fill cells with expected 5-column structure:
        # Col 0: Désignation
        # Col 1: Description
        # Col 2: Qté
        # Col 3: Prix unitaire
        # Col 4: Prix total

        if len(row.cells) < 5:
            logger.warning(f"Table has only {len(row.cells)} columns, expected 5")

        # Clear any existing content from copied row
        for cell in row.cells:
            cell.text = ""

        # Colonne 0: Désignation
        row.cells[0].text = designation
        for paragraph in row.cells[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Colonne 1: Description
        if len(row.cells) > 1:
            row.cells[1].text = description
            for paragraph in row.cells[1].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Colonne 2: Quantité
        if len(row.cells) > 2:
            row.cells[2].text = str(quantity)
            for paragraph in row.cells[2].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Colonne 3: Prix Unitaire HT
        if len(row.cells) > 3:
            row.cells[3].text = f"{unit_price:.2f} €"
            for paragraph in row.cells[3].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Colonne 4: Prix Total HT
        if len(row.cells) > 4:
            row.cells[4].text = f"{total_price:.2f} €"
            for paragraph in row.cells[4].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        logger.info(f"Added offer line: {designation} (Qty: {quantity}, Total: {total_price:.2f}€)")

        # Update table total
        update_table_total(table)

        # Save document
        output_bytes_io = io.BytesIO()
        doc.save(output_bytes_io)
        output_bytes = output_bytes_io.getvalue()

        # Upload to Blob Storage
        blob_url = blob_client.upload_blob(
            container_name=container_name,
            blob_name=working_file_path,
            data=output_bytes,
            overwrite=True
        )

        logger.info(f"Working file updated: {working_file_path}")

        # Return success response
        response_data = {
            "success": True,
            "working_file": working_file_path,
            "service_name": service_name,
            "table_created_or_found": "created" if table_idx == len(doc.tables) - 1 else "found",
            "rows_in_table": len(table.rows),
            "table_total_ht": f"{sum([float(table.rows[i].cells[4].text.replace('€', '').replace(',', '.').strip()) for i in range(1, len(table.rows) - 1)]):.2f} €"
        }

        return func.HttpResponse(
            json.dumps(response_data),
            status_code=200,
            mimetype="application/json"
        )

    except Exception as e:
        logger.error(f"Error adding offer line: {str(e)}", exc_info=True)
        return func.HttpResponse(
            json.dumps({
                "error": "Failed to add offer line",
                "message": str(e)
            }),
            status_code=500,
            mimetype="application/json"
        )
