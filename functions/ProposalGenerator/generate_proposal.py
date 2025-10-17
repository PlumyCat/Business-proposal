"""
Generate proposal documents (Word + PDF) from template and content
"""

import json
import io
import logging
import uuid
from datetime import datetime
from typing import Dict, List, Any, Optional
import azure.functions as func
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import shared modules
import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))
from shared.blob_client import get_blob_client
from shared.dataverse_client import get_dataverse_client
from shared.validators import validate_required_fields
from shared.logger import setup_logger

logger = setup_logger(__name__)


def merge_content_into_template(
    template_doc: Document,
    cleaned_content: Dict[str, Any],
    offers: List[Dict[str, Any]],
    client_info: Dict[str, str]
) -> Document:
    """
    Merge cleaned content and selected offers into template document

    Args:
        template_doc: Loaded Word template document
        cleaned_content: Cleaned content from DocumentProcessor
        offers: List of selected offers from Dataverse
        client_info: Client information (name, contact, etc.)

    Returns:
        Modified Document object
    """
    logger.info("Merging content into template")

    # Replace placeholders in template
    placeholders = {
        "{{CLIENT_NAME}}": client_info.get("name", ""),
        "{{CLIENT_CONTACT}}": client_info.get("contact", ""),
        "{{CLIENT_EMAIL}}": client_info.get("email", ""),
        "{{DATE}}": datetime.now().strftime("%d/%m/%Y"),
        "{{PROPOSAL_NUMBER}}": f"PROP-{datetime.now().strftime('%Y%m%d')}-{str(uuid.uuid4())[:8].upper()}"
    }

    # Replace placeholders in paragraphs
    for paragraph in template_doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    # Replace placeholders in tables
    for table in template_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in placeholders.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)

    # Add cleaned content paragraphs
    if cleaned_content.get("paragraphs"):
        template_doc.add_heading("Contenu du devis", level=1)
        for para in cleaned_content["paragraphs"]:
            if para.get("text") and para["text"].strip():
                p = template_doc.add_paragraph(para["text"])
                if para.get("style") == "Heading1":
                    p.style = 'Heading 1'
                elif para.get("style") == "Heading2":
                    p.style = 'Heading 2'

    # Add cleaned tables
    if cleaned_content.get("tables"):
        template_doc.add_heading("Tables extraites", level=1)
        for table_data in cleaned_content["tables"]:
            if table_data.get("headers") and table_data.get("rows"):
                # Create table with headers + rows
                num_cols = len(table_data["headers"])
                num_rows = len(table_data["rows"]) + 1  # +1 for header row
                table = template_doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Light Grid Accent 1'

                # Add headers
                header_row = table.rows[0]
                for idx, header in enumerate(table_data["headers"]):
                    header_row.cells[idx].text = header
                    header_row.cells[idx].paragraphs[0].runs[0].bold = True

                # Add data rows
                for row_idx, row_data in enumerate(table_data["rows"], start=1):
                    for col_idx, cell_value in enumerate(row_data):
                        table.rows[row_idx].cells[col_idx].text = str(cell_value)

    # Add selected offers
    if offers:
        template_doc.add_page_break()
        template_doc.add_heading("Offres sélectionnées", level=1)

        for offer in offers:
            # Add offer title
            template_doc.add_heading(offer.get("cr_name", "Offre sans nom"), level=2)

            # Add offer description
            if offer.get("cr_description"):
                template_doc.add_paragraph(offer["cr_description"])

            # Add offer details table
            details_table = template_doc.add_table(rows=4, cols=2)
            details_table.style = 'Light List Accent 1'

            details_table.rows[0].cells[0].text = "Catégorie"
            details_table.rows[0].cells[1].text = offer.get("cr_category", "N/A")

            details_table.rows[1].cells[0].text = "Prix unitaire"
            details_table.rows[1].cells[1].text = f"{offer.get('cr_unit_price', 0):.2f} €"

            details_table.rows[2].cells[0].text = "Unité"
            details_table.rows[2].cells[1].text = offer.get("cr_unit", "N/A")

            details_table.rows[3].cells[0].text = "Référence"
            details_table.rows[3].cells[1].text = offer.get("cr_reference", "N/A")

            template_doc.add_paragraph()  # Spacing

    logger.info("Content merge completed")
    return template_doc


def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """
    Convert Word document to PDF

    Note: This is a placeholder implementation.
    Production should use:
    - LibreOffice headless (requires Linux container)
    - Azure Logic Apps with Word connector
    - Third-party API (CloudConvert, etc.)

    Args:
        docx_bytes: Word document as bytes

    Returns:
        PDF document as bytes
    """
    logger.warning("PDF conversion not implemented - returning empty bytes")
    logger.info("For production, implement one of: LibreOffice headless, Azure Logic Apps, or CloudConvert API")

    # TODO: Implement actual PDF conversion
    # For now, return empty bytes as placeholder
    return b""


def generate_proposal(req: func.HttpRequest) -> func.HttpResponse:
    """
    Generate proposal documents (Word + PDF)

    Request body:
    {
        "cleaned_content": {...},  // From DocumentProcessor
        "offer_ids": ["guid1", "guid2"],
        "template_id": "guid",  // Optional, uses default if not provided
        "client_info": {
            "name": "Client Name",
            "contact": "John Doe",
            "email": "john@example.com"
        }
    }

    Response:
    {
        "success": true,
        "word_url": "https://...",
        "pdf_url": "https://...",
        "word_file_id": "guid",
        "pdf_file_id": "guid",
        "proposal_number": "PROP-20250117-ABC123"
    }
    """
    logger.info("Generate proposal endpoint called")

    try:
        # Parse request body
        try:
            req_body = req.get_json()
        except ValueError:
            return func.HttpResponse(
                json.dumps({"error": "Invalid JSON in request body"}),
                status_code=400,
                mimetype="application/json"
            )

        # Validate required fields
        required_fields = ["cleaned_content", "offer_ids", "client_info"]
        if not validate_required_fields(req_body, required_fields):
            return func.HttpResponse(
                json.dumps({
                    "error": "Missing required fields",
                    "required": required_fields
                }),
                status_code=400,
                mimetype="application/json"
            )

        cleaned_content = req_body["cleaned_content"]
        offer_ids = req_body["offer_ids"]
        template_id = req_body.get("template_id")
        client_info = req_body["client_info"]

        # Validate client_info fields
        client_required = ["name", "contact", "email"]
        if not validate_required_fields(client_info, client_required):
            return func.HttpResponse(
                json.dumps({
                    "error": "Missing required client_info fields",
                    "required": client_required
                }),
                status_code=400,
                mimetype="application/json"
            )

        # Initialize clients
        blob_client = get_blob_client()
        dataverse_client = get_dataverse_client()

        # Get selected offers from Dataverse
        logger.info(f"Fetching {len(offer_ids)} offers from Dataverse")
        offers = []
        for offer_id in offer_ids:
            try:
                offer = dataverse_client.get_record(
                    entity_set="cr_offres",
                    record_id=offer_id,
                    select=["cr_name", "cr_description", "cr_category", "cr_unit_price", "cr_unit", "cr_reference"]
                )
                offers.append(offer)
            except Exception as e:
                logger.warning(f"Could not fetch offer {offer_id}: {str(e)}")

        if not offers:
            return func.HttpResponse(
                json.dumps({"error": "No valid offers found"}),
                status_code=404,
                mimetype="application/json"
            )

        # Load template from Blob Storage
        template_blob_name = template_id if template_id else "default_template.docx"
        logger.info(f"Loading template: {template_blob_name}")

        try:
            template_bytes = blob_client.download_blob("templates", template_blob_name)
            template_doc = Document(io.BytesIO(template_bytes))
        except Exception as e:
            logger.error(f"Could not load template: {str(e)}")
            return func.HttpResponse(
                json.dumps({
                    "error": "Template not found or could not be loaded",
                    "template": template_blob_name
                }),
                status_code=404,
                mimetype="application/json"
            )

        # Merge content into template
        merged_doc = merge_content_into_template(
            template_doc=template_doc,
            cleaned_content=cleaned_content,
            offers=offers,
            client_info=client_info
        )

        # Generate proposal number
        proposal_number = f"PROP-{datetime.now().strftime('%Y%m%d')}-{str(uuid.uuid4())[:8].upper()}"

        # Save Word document to bytes
        word_bytes_io = io.BytesIO()
        merged_doc.save(word_bytes_io)
        word_bytes = word_bytes_io.getvalue()

        # Upload Word document to Blob Storage
        word_file_id = str(uuid.uuid4())
        word_blob_name = f"{proposal_number}.docx"
        logger.info(f"Uploading Word document: {word_blob_name}")

        word_url = blob_client.upload_blob(
            container_name="propositions-generated",
            blob_name=word_blob_name,
            data=word_bytes,
            overwrite=True
        )

        # Convert to PDF and upload
        pdf_bytes = convert_docx_to_pdf(word_bytes)
        pdf_file_id = str(uuid.uuid4())
        pdf_blob_name = f"{proposal_number}.pdf"

        if pdf_bytes:
            logger.info(f"Uploading PDF document: {pdf_blob_name}")
            pdf_url = blob_client.upload_blob(
                container_name="propositions-generated",
                blob_name=pdf_blob_name,
                data=pdf_bytes,
                overwrite=True
            )
        else:
            logger.warning("PDF conversion returned empty bytes - skipping PDF upload")
            pdf_url = None

        # Return success response
        response_data = {
            "success": True,
            "proposal_number": proposal_number,
            "word_url": word_url,
            "word_file_id": word_file_id,
            "word_blob_name": word_blob_name,
            "pdf_url": pdf_url if pdf_url else None,
            "pdf_file_id": pdf_file_id if pdf_url else None,
            "pdf_blob_name": pdf_blob_name if pdf_url else None,
            "offers_included": len(offers)
        }

        logger.info(f"Proposal generated successfully: {proposal_number}")
        return func.HttpResponse(
            json.dumps(response_data),
            status_code=200,
            mimetype="application/json"
        )

    except Exception as e:
        logger.error(f"Error generating proposal: {str(e)}", exc_info=True)
        return func.HttpResponse(
            json.dumps({
                "error": "Failed to generate proposal",
                "message": str(e)
            }),
            status_code=500,
            mimetype="application/json"
        )
