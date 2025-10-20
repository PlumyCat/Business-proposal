"""
Set Customer Success Info - Met à jour les infos customer success dans le document
"""

import json
import logging
import io
import os
from typing import Dict
import azure.functions as func
from docx import Document

import sys
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))
from shared.blob_client import get_blob_client
from shared.validators import validate_required_fields
from shared.logger import setup_logger
from shared.config import CONTAINER_TEMPLATES, get_user_file_path

logger = setup_logger(__name__)


def replace_customer_placeholders(doc: Document, customer_info: Dict[str, str]) -> Document:
    """
    Remplace les placeholders customer success dans le document

    Placeholders:
    - {{CS_NAME}} → nom customer success
    - {{CS_TEL}} → téléphone
    - {{CS_EMAIL}} → email

    Args:
        doc: Document Word
        customer_info: Dict avec name, tel, email

    Returns:
        Document modifié
    """
    placeholders = {
        "{{CS_NAME}}": customer_info.get("name", ""),
        "{{CS_TEL}}": customer_info.get("tel", ""),
        "{{CS_EMAIL}}": customer_info.get("email", "")
    }

    # Remplacer dans les paragraphes
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                # Remplacer dans les runs pour conserver le formatage
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
                logger.debug(f"Replaced {placeholder} in paragraph")

    # Remplacer dans les tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in placeholders.items():
                        if placeholder in paragraph.text:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, value)
                            logger.debug(f"Replaced {placeholder} in table cell")

    logger.info("Customer success placeholders replaced")
    return doc


def set_customer_info(req: func.HttpRequest) -> func.HttpResponse:
    """
    Met à jour les infos customer success dans le document de travail

    Le bot peut récupérer automatiquement ces infos et appeler cet endpoint
    pour les insérer dans le document en une seule fois.

    Request body:
    {
        "user_folder": "Eric FER",
        "customer_success": {
            "name": "Eric FER",
            "tel": "06 12 34 56 78",
            "email": "eric.fer@company.com"
        }
    }

    Response:
    {
        "success": true,
        "working_file": "Eric FER/temp_working.docx",
        "customer_success": {...}
    }
    """
    logger.info("Set customer info endpoint called")

    try:
        # Parse request body
        try:
            req_body = req.get_json()
        except ValueError:
            return func.HttpResponse(
                json.dumps({"error": "Invalid JSON in request body"}),
                status_code=400,
                mimetype="application/json"
            )

        # Validate required fields
        if not validate_required_fields(req_body, ["user_folder", "customer_success"]):
            return func.HttpResponse(
                json.dumps({
                    "error": "Missing required fields",
                    "required": ["user_folder", "customer_success"]
                }),
                status_code=400,
                mimetype="application/json"
            )

        user_folder = req_body["user_folder"]
        customer_success = req_body["customer_success"]

        # Validate customer_success fields
        if not validate_required_fields(customer_success, ["name", "tel", "email"]):
            return func.HttpResponse(
                json.dumps({
                    "error": "Missing customer_success fields",
                    "required": ["name", "tel", "email"]
                }),
                status_code=400,
                mimetype="application/json"
            )

        logger.info(f"Setting customer info for {user_folder}")

        # Initialize Blob client
        blob_client = get_blob_client()
        container_name = CONTAINER_TEMPLATES

        # Load working file
        working_file_path = get_user_file_path(user_folder, "temp_working.docx")

        try:
            working_bytes = blob_client.download_blob(container_name, working_file_path)
            logger.info(f"Downloaded working file: {working_file_path}")
        except Exception as e:
            logger.error(f"Failed to download working file: {str(e)}")
            return func.HttpResponse(
                json.dumps({
                    "error": "Working file not found",
                    "working_file": working_file_path,
                    "hint": "Call /api/proposal/prepare-template first"
                }),
                status_code=404,
                mimetype="application/json"
            )

        # Load document
        doc = Document(io.BytesIO(working_bytes))

        # Replace placeholders
        doc = replace_customer_placeholders(doc, customer_success)

        # Save document
        output_bytes_io = io.BytesIO()
        doc.save(output_bytes_io)
        output_bytes = output_bytes_io.getvalue()

        # Upload to Blob Storage
        blob_url = blob_client.upload_blob(
            container_name=container_name,
            blob_name=working_file_path,
            data=output_bytes,
            overwrite=True
        )

        logger.info(f"Working file updated with customer info: {working_file_path}")

        # Return success response
        response_data = {
            "success": True,
            "working_file": working_file_path,
            "customer_success": customer_success
        }

        return func.HttpResponse(
            json.dumps(response_data),
            status_code=200,
            mimetype="application/json"
        )

    except Exception as e:
        logger.error(f"Error setting customer info: {str(e)}", exc_info=True)
        return func.HttpResponse(
            json.dumps({
                "error": "Failed to set customer info",
                "message": str(e)
            }),
            status_code=500,
            mimetype="application/json"
        )
