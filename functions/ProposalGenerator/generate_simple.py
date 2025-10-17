"""
Generate Proposal - Version simplifiée
Charge fichier de travail + ajoute offres + génère Word/PDF
"""

import json
import logging
import io
import os
from datetime import datetime
from typing import Dict, List
import azure.functions as func
from docx import Document

import sys
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))
from shared.blob_client import get_blob_client
from shared.dataverse_client import get_dataverse_client
from shared.sharepoint_client import get_sharepoint_client
from shared.validators import validate_required_fields
from shared.logger import setup_logger

logger = setup_logger(__name__)


def add_offers_to_table(doc: Document, offers: List[Dict]) -> int:
    """
    Ajoute les offres sélectionnées dans le premier tableau du document

    Suppose que le tableau a déjà des headers et que les colonnes sont:
    [Description, Quantité, Prix Unitaire, Prix Total] ou similaire

    Args:
        doc: Document Word
        offers: Liste des offres depuis Dataverse

    Returns:
        Nombre de lignes ajoutées
    """
    if not doc.tables:
        logger.warning("No tables found in document")
        return 0

    # Utiliser le premier tableau (on peut améliorer plus tard)
    table = doc.tables[0]

    rows_added = 0
    for offer in offers:
        # Ajouter une nouvelle ligne
        row = table.add_row()

        # Remplir les cellules
        # Colonnes: Description | Quantité | Prix Unitaire | Prix Total
        try:
            # Colonne 0: Description (nom + description de l'offre)
            description = offer.get("cr_name", "")
            if offer.get("cr_description"):
                description += f"\n{offer['cr_description']}"
            row.cells[0].text = description

            # Colonne 1: Quantité (à définir, par défaut 1)
            if len(row.cells) > 1:
                row.cells[1].text = "1"

            # Colonne 2: Prix Unitaire
            if len(row.cells) > 2:
                unit_price = offer.get("cr_unit_price", 0)
                unit = offer.get("cr_unit", "")
                row.cells[2].text = f"{unit_price:.2f} € / {unit}"

            # Colonne 3: Prix Total (quantité * prix unitaire, par défaut quantité=1)
            if len(row.cells) > 3:
                row.cells[3].text = f"{unit_price:.2f} €"

            rows_added += 1
            logger.debug(f"Added offer: {offer.get('cr_name')}")

        except Exception as e:
            logger.error(f"Error adding offer to table: {str(e)}")

    logger.info(f"Added {rows_added} offers to table")
    return rows_added


def generate_proposal_simple(req: func.HttpRequest) -> func.HttpResponse:
    """
    Génère une proposition à partir du fichier de travail + offres sélectionnées

    Workflow:
    1. Charge temp_working.docx depuis Blob Storage
    2. Récupère les offres sélectionnées depuis Dataverse
    3. Ajoute les offres dans le tableau
    4. Sauvegarde dans users/{user_folder}/proposition_{date}.docx
    5. Convertit en PDF via SharePoint
    6. Retourne les URLs de téléchargement

    Request body:
    {
        "working_file": "users/Jean Dupont/temp_working.docx",
        "offer_ids": ["guid1", "guid2", "guid3"],
        "user_folder": "Jean Dupont"
    }

    Response:
    {
        "success": true,
        "word_file": "users/Jean Dupont/proposition_20251017_1430.docx",
        "word_url": "https://...",
        "pdf_file": "users/Jean Dupont/proposition_20251017_1430.pdf",
        "pdf_url": "https://...",
        "offers_added": 3
    }
    """
    logger.info("Generate proposal (simple) endpoint called")

    try:
        # Parse request body
        try:
            req_body = req.get_json()
        except ValueError:
            return func.HttpResponse(
                json.dumps({"error": "Invalid JSON in request body"}),
                status_code=400,
                mimetype="application/json"
            )

        # Validate required fields
        if not validate_required_fields(req_body, ["working_file", "offer_ids", "user_folder"]):
            return func.HttpResponse(
                json.dumps({
                    "error": "Missing required fields",
                    "required": ["working_file", "offer_ids", "user_folder"]
                }),
                status_code=400,
                mimetype="application/json"
            )

        working_file = req_body["working_file"]
        offer_ids = req_body["offer_ids"]
        user_folder = req_body["user_folder"]

        if not offer_ids or not isinstance(offer_ids, list):
            return func.HttpResponse(
                json.dumps({"error": "offer_ids must be a non-empty array"}),
                status_code=400,
                mimetype="application/json"
            )

        logger.info(f"Generating proposal for {user_folder} with {len(offer_ids)} offers")

        # Initialize clients
        blob_client = get_blob_client()
        dataverse_client = get_dataverse_client()
        container_name = os.environ.get("BLOB_CONTAINER_DEVIS", "devis-sources")

        # 1. Charger le fichier de travail
        try:
            working_bytes = blob_client.download_blob(container_name, working_file)
            logger.info(f"Downloaded working file: {working_file}")
        except Exception as e:
            logger.error(f"Failed to download working file: {str(e)}")
            return func.HttpResponse(
                json.dumps({
                    "error": "Working file not found",
                    "working_file": working_file
                }),
                status_code=404,
                mimetype="application/json"
            )

        doc = Document(io.BytesIO(working_bytes))

        # 2. Récupérer les offres depuis Dataverse
        logger.info(f"Fetching {len(offer_ids)} offers from Dataverse")
        offers = []
        for offer_id in offer_ids:
            try:
                offer = dataverse_client.get_record(
                    entity_set="cr_offres",
                    record_id=offer_id,
                    select=["cr_name", "cr_description", "cr_category", "cr_unit_price", "cr_unit", "cr_reference"]
                )
                offers.append(offer)
            except Exception as e:
                logger.warning(f"Could not fetch offer {offer_id}: {str(e)}")

        if not offers:
            return func.HttpResponse(
                json.dumps({"error": "No valid offers found"}),
                status_code=404,
                mimetype="application/json"
            )

        # 3. Ajouter les offres dans le tableau
        rows_added = add_offers_to_table(doc, offers)

        # 4. Générer nom de fichier avec timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        proposal_filename = f"proposition_{timestamp}"

        # Sauvegarder le document Word
        word_bytes_io = io.BytesIO()
        doc.save(word_bytes_io)
        word_bytes = word_bytes_io.getvalue()

        # Upload Word vers Blob Storage
        word_file_path = f"users/{user_folder}/{proposal_filename}.docx"
        word_url = blob_client.upload_blob(
            container_name=container_name,
            blob_name=word_file_path,
            data=word_bytes,
            overwrite=True
        )
        logger.info(f"Word file saved to: {word_file_path}")

        # 5. Convertir en PDF via SharePoint
        pdf_url = None
        pdf_file_path = None

        try:
            sharepoint_client = get_sharepoint_client()
            pdf_bytes = sharepoint_client.convert_word_to_pdf(
                word_content=word_bytes,
                file_name=f"temp_{proposal_filename}.docx"
            )

            if pdf_bytes:
                # Upload PDF vers Blob Storage
                pdf_file_path = f"users/{user_folder}/{proposal_filename}.pdf"
                pdf_url = blob_client.upload_blob(
                    container_name=container_name,
                    blob_name=pdf_file_path,
                    data=pdf_bytes,
                    overwrite=True
                )
                logger.info(f"PDF file saved to: {pdf_file_path}")
            else:
                logger.warning("PDF conversion returned empty bytes")

        except Exception as e:
            logger.error(f"PDF conversion failed: {str(e)}")
            # Continue sans PDF, pas bloquant

        # 6. Retourner résultat
        response_data = {
            "success": True,
            "word_file": word_file_path,
            "word_url": word_url,
            "pdf_file": pdf_file_path if pdf_url else None,
            "pdf_url": pdf_url if pdf_url else None,
            "offers_added": rows_added
        }

        logger.info(f"Proposal generated successfully for {user_folder}")
        return func.HttpResponse(
            json.dumps(response_data),
            status_code=200,
            mimetype="application/json"
        )

    except Exception as e:
        logger.error(f"Error generating proposal: {str(e)}", exc_info=True)
        return func.HttpResponse(
            json.dumps({
                "error": "Failed to generate proposal",
                "message": str(e)
            }),
            status_code=500,
            mimetype="application/json"
        )
