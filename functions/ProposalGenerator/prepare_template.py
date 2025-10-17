"""
Prepare Template - Charge un template général et remplace les infos customer success
"""

import json
import logging
import io
import os
from typing import Dict
import azure.functions as func
from docx import Document

import sys
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))
from shared.blob_client import get_blob_client
from shared.validators import validate_required_fields
from shared.logger import setup_logger
from shared.config import CONTAINER_TEMPLATES, get_template_path, get_user_file_path

logger = setup_logger(__name__)


def replace_customer_success_placeholders(doc: Document, customer_success: Dict[str, str]) -> Document:
    """
    Remplace les placeholders customer success dans le document

    Placeholders:
    - {{CS_NAME}} → nom customer success
    - {{CS_TEL}} → téléphone
    - {{CS_EMAIL}} → email

    Args:
        doc: Document Word
        customer_success: Dict avec name, tel, email

    Returns:
        Document modifié
    """
    placeholders = {
        "{{CS_NAME}}": customer_success.get("name", ""),
        "{{CS_TEL}}": customer_success.get("tel", ""),
        "{{CS_EMAIL}}": customer_success.get("email", "")
    }

    # Remplacer dans les paragraphes
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
                logger.debug(f"Replaced {placeholder} in paragraph")

    # Remplacer dans les tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in placeholders.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)
                        logger.debug(f"Replaced {placeholder} in table cell")

    logger.info("Customer success placeholders replaced")
    return doc


def prepare_template(req: func.HttpRequest) -> func.HttpResponse:
    """
    Prépare un template général avec les infos customer success

    Workflow:
    1. Charge le template depuis general/{template_name}
    2. Remplace {{CS_NAME}}, {{CS_TEL}}, {{CS_EMAIL}}
    3. Sauvegarde dans users/{user_folder}/temp_working.docx

    Request body:
    {
        "template_name": "template_standard.docx",
        "customer_success": {
            "name": "Jean Dupont",
            "tel": "06 12 34 56 78",
            "email": "jean.dupont@company.com"
        },
        "user_folder": "Marie Martin"
    }

    Response:
    {
        "success": true,
        "working_file": "users/Marie Martin/temp_working.docx",
        "blob_url": "https://...",
        "template_used": "template_standard.docx"
    }
    """
    logger.info("Prepare template endpoint called")

    try:
        # Parse request body
        try:
            req_body = req.get_json()
        except ValueError:
            return func.HttpResponse(
                json.dumps({"error": "Invalid JSON in request body"}),
                status_code=400,
                mimetype="application/json"
            )

        # Validate required fields
        if not validate_required_fields(req_body, ["template_name", "customer_success", "user_folder"]):
            return func.HttpResponse(
                json.dumps({
                    "error": "Missing required fields",
                    "required": ["template_name", "customer_success", "user_folder"]
                }),
                status_code=400,
                mimetype="application/json"
            )

        template_name = req_body["template_name"]
        customer_success = req_body["customer_success"]
        user_folder = req_body["user_folder"]

        # Validate customer_success fields
        if not validate_required_fields(customer_success, ["name", "tel", "email"]):
            return func.HttpResponse(
                json.dumps({
                    "error": "Missing customer_success fields",
                    "required": ["name", "tel", "email"]
                }),
                status_code=400,
                mimetype="application/json"
            )

        logger.info(f"Preparing template: {template_name} for user: {user_folder}")

        # Initialize Blob client
        blob_client = get_blob_client()
        container_name = CONTAINER_TEMPLATES

        # Charger le template depuis general/
        template_path = get_template_path(template_name)

        try:
            template_bytes = blob_client.download_blob(container_name, template_path)
            logger.info(f"Downloaded template: {template_path}")
        except Exception as e:
            logger.error(f"Failed to download template: {str(e)}")
            return func.HttpResponse(
                json.dumps({
                    "error": "Template not found",
                    "template_path": template_path
                }),
                status_code=404,
                mimetype="application/json"
            )

        # Charger le document Word
        doc = Document(io.BytesIO(template_bytes))

        # Remplacer les placeholders customer success
        doc = replace_customer_success_placeholders(doc, customer_success)

        # Sauvegarder le document préparé
        output_bytes_io = io.BytesIO()
        doc.save(output_bytes_io)
        output_bytes = output_bytes_io.getvalue()

        # Upload vers Blob Storage comme fichier de travail
        working_file_path = get_user_file_path(user_folder, "temp_working.docx")

        blob_url = blob_client.upload_blob(
            container_name=container_name,
            blob_name=working_file_path,
            data=output_bytes,
            overwrite=True
        )

        logger.info(f"Working file saved to: {working_file_path}")

        # Retourner résultat
        response_data = {
            "success": True,
            "working_file": working_file_path,
            "blob_url": blob_url,
            "template_used": template_name
        }

        return func.HttpResponse(
            json.dumps(response_data),
            status_code=200,
            mimetype="application/json"
        )

    except Exception as e:
        logger.error(f"Error preparing template: {str(e)}", exc_info=True)
        return func.HttpResponse(
            json.dumps({
                "error": "Failed to prepare template",
                "message": str(e)
            }),
            status_code=500,
            mimetype="application/json"
        )
